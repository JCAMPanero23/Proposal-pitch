"""
Generate financial-model-revision-150m-v2.docx — supplementary calculation
memo for the Canvas From The Sky investor pitch deck (PitchDeck_v2 index.html).

The `-v2` suffix matches the PitchDeck v2 deck — it is a filename version,
not a financial-model revision number.

This document is read alongside the deck by company higher-ups. It documents
every assumption change vs. pitch-blueprint.md and shows the arithmetic
behind each headline number. No CEO approval questions — the model is the
output, not the input.

⚠️  WARNING — re-running this script OVERWRITES financial-model-revision-150m-v2.docx
    with a fresh auto-generated layout. The committed docx contains MANUAL
    layout polish done in Word after this script last ran. If you need to
    update the content, edit the source data in this file, run the script,
    then re-apply any manual layout work in Word and re-export the PDF.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GOLD = RGBColor(0x8C, 0x6F, 0x2A)
TEXT = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)


def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_para(doc, text, *, bold=False, italic=False, size=11, color=TEXT,
             align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = TEXT
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = GOLD
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = TEXT
    return p


def add_table(doc, header, rows, *, col_widths=None, first_col_left=True,
              section_marker='__SECTION__'):
    """Render a table. Rows whose first cell starts with `section_marker`
    are rendered as a merged, shaded sub-header band spanning all columns."""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    # Header
    for i, txt in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, 'F1E9D2')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(txt))
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = TEXT
    # Body
    for r, row in enumerate(rows, start=1):
        first = str(row[0])
        if first.startswith(section_marker):
            label = first[len(section_marker):].strip()
            cells = table.rows[r].cells
            merged = cells[0]
            for c in cells[1:]:
                merged = merged.merge(c)
            set_cell_shading(merged, 'EFE2C2')
            merged.text = ''
            p = merged.paragraphs[0]
            run = p.add_run(label)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT
            continue
        for i, val in enumerate(row):
            cell = table.rows[r].cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT
            if i == 0 and first_col_left:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif i > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return table


def bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT
    return p


# ──────────────────────────────────────────────────────────────────────
# Build the document
# ──────────────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── COVER ────────────────────────────────────────────────────────────
add_para(doc, 'FINANCIAL MODEL REVISION MEMO',
         bold=True, size=10, color=MUTED)
add_para(doc, 'Canvas From The Sky',
         bold=True, size=26, color=TEXT)
add_para(doc, 'Supplementary calculation document for the investor pitch deck',
         italic=True, size=12, color=MUTED)
add_para(doc, '')
add_para(doc, '150 m peak operating height · 25-guest commercial cap · '
              '16-minute day cycle · 20-minute per-batch night cycle · '
              'optional 10-min LED-field dwell · baseline ARPU retained',
         size=11, color=TEXT)
add_para(doc, '')

meta = doc.add_table(rows=4, cols=2)
meta.autofit = False
for i, w in enumerate([3.5, 12.5]):
    for cell in meta.columns[i].cells:
        cell.width = Cm(w)
meta_rows = [
    ('Audience', 'Company executive review — read alongside PitchDeck v2 (index.html)'),
    ('Date',     '15 May 2026'),
    ('Status',   'Aligned with the revised PitchDeck v2 HTML'),
    ('Sources',  'pitch-blueprint.md (baseline) · financial-model-revision-150m.md (working draft)'),
]
for i, (k, v) in enumerate(meta_rows):
    c0, c1 = meta.rows[i].cells
    c0.text = ''
    run = c0.paragraphs[0].add_run(k)
    run.bold = True; run.font.name = 'Calibri'; run.font.size = Pt(10); run.font.color.rgb = MUTED
    c1.text = ''
    run = c1.paragraphs[0].add_run(v)
    run.font.name = 'Calibri'; run.font.size = Pt(10); run.font.color.rgb = TEXT

# ── 1. PURPOSE ───────────────────────────────────────────────────────
h1(doc, '1. Purpose of this memo')
add_para(doc,
    'The investor pitch deck (PitchDeck v2, index.html) communicates the '
    'revised Canvas From The Sky financial model in summary form. This '
    'memo is the calculation backing it. Each headline figure in the deck '
    'can be traced here to the operating assumption it rests on, and each '
    'operating assumption is shown against the previous pitch-blueprint.md '
    'baseline so executives can see what changed and why.')
add_para(doc,
    'The revision moves five operating assumptions versus the '
    'pitch-blueprint.md baseline:')
bullet(doc, 'Maximum balloon height: 300 m  →  150 m (true new peak — the attraction is no longer marketed as 300 m capable).')
bullet(doc, 'Commercial basket loading: 30 guests  →  25 guests (the physical gondola remains 30-passenger).')
bullet(doc, 'Day cycle: 30 min · 6 rides/day  →  16 min · 11 rides/day. The 150 m peak shortens vertical motion from ~13.5 min to ~6.75 min round-trip, freeing the cycle while preserving a generous 5-min panoramic peak hold.')
bullet(doc, 'Night cycle: a single 41-minute cycle  →  a 20-minute per-batch cadence, with an on-platform photo dwell after the gondola lands and the LED entry hallway closed during the main show.')
bullet(doc, 'Night LED-field dwell: guests may linger inside the LED ring for up to 10 minutes after disembarking, turning the field itself into a photo destination. No throughput impact at the gondola; modest F&B uplift.')

# ── 2. AT A GLANCE — CONSOLIDATED SNAPSHOT ───────────────────────────
h1(doc, '2. At a glance — baseline vs. revised')
add_para(doc,
    'A single-table view of every assumption and outcome that moved between '
    'the pitch-blueprint.md baseline and the revised model in PitchDeck v2. '
    'Sections §3 onward show the calculations behind each row.')
add_table(doc,
    header=['Dimension', 'Baseline (blueprint)', 'Revised (this memo)', 'Change'],
    rows=[
        ['__SECTION__Operating assumptions'],
        ['Peak balloon altitude',          '300 m',                  '150 m',                  'True new peak'],
        ['Gondola commercial capacity',    '30 guests / batch',      '25 guests / batch',      '– 5 guests'],
        ['Operating days (day flights)',   '300 / year',             '250 / year',             '– 50 days'],
        ['Operating nights (LED show)',    '300 / year',             '250 / year',             '– 50 nights'],
        ['Day cycle',                      '30 min · 6 rides/day',   '16 min · 11 rides/day',  '+ 5 rides/day'],
        ['Night cycle',                    '41 min · 8 rides/night', '20 min · 16 batches/night', '+ 8 batches/night'],
        ['Night LED-field dwell',          '—',                      'Up to 10 min optional',  'New UX layer'],
        ['Day blended ARPU',               'AED 170',                'AED 170',                'Unchanged'],
        ['Night blended ARPU',             'AED 361',                'AED 361',                'Unchanged'],
        ['Base case load factor',          '70 %',                   '58 %',                   '– 12 pts'],
        ['__SECTION__Capacity outputs'],
        ['Day seat capacity',              '180 / day',              '275 / day',              '+ 52.8 %'],
        ['Night seat capacity',            '240 / night',            '400 / night',            '+ 66.7 %'],
        ['__SECTION__Financial outputs (Base Case)'],
        ['CAPEX (single-spec aluminum)',   'AED 25.9 M',             'AED 25.9 M',             'Unchanged'],
        ['Annual revenue',                 'AED 24.7 M',             'AED 25.5 M',             '+ AED 0.8 M'],
        ['EBITDA',                         'AED 14.9 M',             'AED 15.6 M',             '+ AED 0.7 M'],
        ['EBITDA margin',                  '60 %',                   '61 %',                   '+ 1 pt'],
        ['Simple payback',                 '~1.6 yrs',               '~1.7 yrs',               '+ 0.1 yrs'],
        ['10-year IRR',                    '~58 %',                  '~45 %',                  '– 13 pts'],
        ['10-year NPV at 8 %',             'AED 85 M',               'AED 73.4 M',             '– AED 11.6 M'],
        ['10-year MOIC',                   '~5.5×',                  '~6.0×',                  '+ 0.5×'],
    ],
    col_widths=[5.0, 4.0, 4.0, 3.0])
add_para(doc, '')
add_para(doc,
    'Read in one line: the revised model is now slightly stronger than the '
    'blueprint at the project level. Two changes do most of the work — the '
    '20-minute night cadence (which more than doubles nightly batch count) '
    'and the 16-minute day cadence (which almost doubles day rides within '
    'the existing 3-hour daytime window). Both are enabled by the 150 m '
    'altitude shortening vertical motion. The 250-day operating year and '
    'the 25-guest cap remain in place as the conservative envelope.',
    italic=True, color=MUTED, size=10)

# ── 3. BASELINE FROM pitch-blueprint.md ──────────────────────────────
h1(doc, '3. Baseline reference — pitch-blueprint.md')
add_para(doc,
    'The previous investor model documented in pitch-blueprint.md (Part 1, '
    'Part 7, Part 8) used the following operating assumptions and outputs. '
    'These are the figures the revision is measured against.')
h3(doc, 'Operating assumptions (baseline)')
add_table(doc,
    header=['Parameter', 'Baseline value', 'Source in blueprint'],
    rows=[
        ['Maximum balloon height',         '300 m',                    'Part 1 · Appendix A'],
        ['Gondola commercial capacity',    '30 guests',                'Part 1 · Appendix A'],
        ['Operating days (day flights)',   '300 / year',               'Part 7'],
        ['Operating nights (LED show)',    '300 / year',               'Part 7'],
        ['Day flight cycle',               '30 min · 6 rides/day',     'Part 1'],
        ['Night LED show cycle',           '41 min · 8 rides/night',   'Part 1'],
        ['Day blended ARPU',               'AED 170',                  'Part 7'],
        ['Night blended ARPU',             'AED 361',                  'Part 7'],
        ['Base case load factor',          '70 %',                     'Part 7'],
        ['All-Access pass usage factor',   '1.4 rides / unique guest', 'Part 7'],
    ],
    col_widths=[6.0, 4.5, 5.5])
h3(doc, 'Baseline financial outputs')
add_table(doc,
    header=['Metric', 'Baseline'],
    rows=[
        ['Base annual revenue',  'AED 24.7 M'],
        ['Strong-case revenue',  'AED 39.7 M'],
        ['Base annual OpEx',     'AED 9.8 M'],
        ['Base EBITDA / margin', 'AED 14.9 M · 60 %'],
        ['Base simple payback',  '~1.6 years'],
        ['10-year IRR',          '~58 %'],
        ['10-year NPV at 8 %',   'AED 85 M'],
        ['10-year MOIC',         '~5.5×'],
    ],
    col_widths=[8.0, 6.0])
add_para(doc,
    'CAPEX in the baseline is AED 25.9 M for the High Tier (aluminum tile) '
    'and AED 25.3 M for the Low Tier (GFRC tile). This memo uses the '
    'AED 25.9 M figure to preserve a single-spec comparison.',
    italic=True, color=MUTED, size=10)

# ── 4. REVISED OPERATING ASSUMPTIONS ─────────────────────────────────
h1(doc, '4. Revised operating assumptions')
add_table(doc,
    header=['Parameter', 'Revised value', 'Rationale'],
    rows=[
        ['Maximum balloon height',
         '150 m',
         'New true peak. Lower operating altitude under typical Saadiyat wind/safety conditions. The product is no longer marketed as 300 m-capable.'],
        ['Gondola commercial capacity',
         '25 guests',
         'More conservative than the 30-pax physical capacity. Improves guest comfort and gives an operating margin for staff and accessibility cases.'],
        ['Operating days (day flights)',
         '250 / year',
         'Higher weather-cancellation and operational-downtime allowance than the baseline 300 days. ~32 % cancellation buffer.'],
        ['Operating nights (LED show)',
         '250 / year',
         'Same buffer as day operations.'],
        ['Day flight structure',
         '11 rides / day (16-min cycle)',
         'Cycle compressed from baseline 30 min to 16 min because the 150 m peak shortens vertical motion from ~13.5 min to ~6.75 min round-trip. Peak panoramic hold retained at 5 min (the panorama IS the day product). Day window stays at 3 hrs (9am – 12pm) to keep the night product clean.'],
        ['Night cycle',
         '20 minutes per batch',
         'Faster cycle becomes feasible because ascent and descent are shorter at 150 m. See section 5 for the per-batch timeline.'],
        ['Night batches per night',
         '16 / night',
         '3 batches per hour × ~5.5-hour night window = 16.5, rounded down for buffer.'],
        ['Night LED-field dwell',
         'Up to 10 min optional',
         'After the gondola lands, guests may linger inside the LED ring for photos. No gondola throughput impact — the next batch still launches every 20 min. Strengthens the All-Access pass narrative (3 themes + dwell time between) and lifts F&B opportunity.'],
        ['All-Access pass usage factor',
         '1.4 rides / unique guest',
         'Unchanged from baseline. The 20-minute cycle actually improves the All-Access proposition (multi-theme viewing is easier in one night).'],
        ['Base case load factor',
         '58 %',
         'Below the baseline 70 % Dubai Balloon benchmark. Acknowledges a slower ramp-up and a more conservative steady-state target.'],
        ['Day blended ARPU',
         'AED 170',
         'Unchanged. See section 7 for the rationale for retaining ARPU at 150 m.'],
        ['Night blended ARPU',
         'AED 361',
         'Unchanged. See section 7.'],
        ['CAPEX (single-spec aluminum)',
         'AED 25.9 M',
         'High-tier aluminum tile assumption from the blueprint.'],
    ],
    col_widths=[4.5, 3.5, 8.0])

# ── 4B. LOADING RATIONALE ────────────────────────────────────────────
h1(doc, '4B. Loading rationale — the 25-guest cap')
add_para(doc,
    'The model uses 25 guests/batch as the commercial cap versus the '
    'baseline pitch-blueprint.md\'s 30 (and versus the original Hero '
    'technical proposal page-9 number of 12). The gondola itself is '
    'unchanged across all three numbers: it is the same certified '
    '30-pax envelope with the same safety margin. The 25-guest figure '
    'is a commercial / operational choice, not a hardware constraint.')
add_para(doc,
    'The defensible framing for investor diligence is that the loading '
    'uplift is a demand-side claim attributable to the LED canvas + '
    'night-show product, not a re-rating of the gondola:')
bullet(doc, 'Original Hero proposal (page 9) used 12 pax/flight × 70% load = 8.4 effective pax/flight. This was likely derived from real Dubai Balloon operating averages — Hero has 250,000+ flights of historical data.')
bullet(doc, 'This model\'s base case uses 25 pax × 58% load = 14.5 effective pax/flight — a +73% uplift in effective per-flight occupancy, or only +21% over the implied Dubai actual of 12 pax/flight average.')
bullet(doc, 'That +21% over Dubai actuals is the claim attributable to the LED canvas + premium night product + museum cross-discount + Saadiyat\'s 10% YoY visitor growth — a demand-side thesis that an investor can validate.')
bullet(doc, 'At no point does the gondola operate beyond its certified safety envelope. 25 of 30 = 83% of physical capacity, leaving margin for staff and accessibility cases.')
add_table(doc,
    header=['Loading line', 'Pax/flight', 'Effective load', 'Effective pax/flight'],
    rows=[
        ['Hero proposal (Dubai actuals)', '12', '70 %',  '8.4'],
        ['Blueprint (envelope max)',      '30', '70 %',  '21.0'],
        ['This model (base case)',        '25', '58 %',  '14.5'],
    ],
    col_widths=[6.5, 3.0, 3.0, 3.5])
add_para(doc,
    'Recommended deck phrasing for investor diligence: "Same gondola, '
    'same certification, same safety envelope. The +21% effective '
    'loading vs. Dubai actuals is the demand uplift attributable to the '
    'LED canvas night product and Saadiyat\'s cultural-district ramp."',
    italic=True, color=MUTED, size=10)

# ── 4C. DAY CYCLE DESIGN AT 150 M ────────────────────────────────────
h1(doc, '4C. Day cycle design at 150 m')
add_para(doc,
    'At 150 m the day cycle compresses from baseline 30 min to ~16 min '
    'because vertical motion drops from 13.5 min (at 300 m) to 6.75 min '
    '(at 150 m). The day product is no LED show — it is the panoramic '
    'view — so cycle compression is not constrained by show length.')
h3(doc, 'Per-flight timeline (T = minutes from flight start)')
add_table(doc,
    header=['Phase', 'Time', 'Duration', 'What happens'],
    rows=[
        ['Boarding',                  '0:00 – 2:00',   '2.0 min',
         'Staff-assisted boarding, 25 guests, safety brief.'],
        ['Ascent to 150 m',           '2:00 – 5:00',   '3.0 min',
         '50 m/min ascent. Museum-and-skyline reveal.'],
        ['Peak panoramic hold',       '5:00 – 10:00',  '5.0 min',
         'Louvre Abu Dhabi · Zayed NM · Guggenheim · coastline · skyline. The money moment for photos.'],
        ['Descent 150 m → 0',         '10:00 – 13:45', '3.75 min',
         '40 m/min descent to the platform.'],
        ['Disembark + turnover',      '13:45 – 16:00', '2.25 min',
         'Next batch begins boarding during disembark.'],
    ],
    col_widths=[3.5, 2.5, 2.0, 8.0])
add_para(doc, '')
h3(doc, 'Why 16 min (and not less)')
bullet(doc, 'A 5-min panoramic peak hold preserves the photo / Instagram moment that the day product is sold on. Anything shorter starts to feel rushed for families and accessibility cases.')
bullet(doc, 'At 15-min cycle the peak shrinks to 4 min — workable, but compromises the experience for marginal throughput gain.')
bullet(doc, 'At 17-min cycle the peak grows to 6 min but throughput drops from 11 to 10 rides/day in the 3-hour window. The day window itself is the binding constraint, not the cycle.')
h3(doc, 'Day window stays at 3 hours (9am – 12pm)')
add_para(doc,
    'The model deliberately keeps the day window short (3 hr) to protect '
    'the night product\'s start. Extending the day window to a full '
    '10-hour day would add ~AED 16 M of annual day revenue at base case '
    'load — this is carried as an upside lever for a later phase, not '
    'in the headline numbers. The current decision: keep daytime '
    'product premium and brief; concentrate ticket revenue in the '
    'evening LED show product.',
    size=10, color=MUTED, italic=True)

# ── 5. NIGHT CYCLE DESIGN AT 150 M ───────────────────────────────────
h1(doc, '5. Night cycle design at 150 m')
add_para(doc,
    'The 20-minute per-batch cadence is the operational hinge of the revised '
    'model. It works because a 150 m peak shortens ascent and descent versus '
    'the 300 m baseline, freeing time for a compressed show window. Two '
    'operating rules protect the experience:')
bullet(doc, 'The LED entry hallway is closed in both directions during the main show. Guests cannot exit through it once the gondola has descended to the 25 m show altitude, so no one casts shadows across the canvas that the gondola is viewing from above.')
bullet(doc, 'After the gondola lands, guests stay on the boarding platform for a short photo dwell. This gives them a soft, memorable exit beat — and it keeps them off the LED hallway until the next batch has cleared it.')
add_para(doc, '')
h3(doc, 'Per-batch timeline (T = minutes from batch start)')
add_table(doc,
    header=['Phase', 'Time', 'Duration', 'LED hallway', 'What happens'],
    rows=[
        ['Boarding from LED hallway', '0:00 – 2:00', '~2.0 min',  'Open (entry)',
         'Guests walk in across the lit hallway, board the gondola, safety brief.'],
        ['Ascent to 150 m',           '2:00 – 5:00', '~3.0 min',  'Open (ambient)',
         '50 m/min ascent. LED field dims under a rising animation.'],
        ['Peak hold at 150 m',        '5:00 – 7:30', '~2.5 min',  'Open (ambient)',
         'Three museums + Gulf coastline + Abu Dhabi skyline in view.'],
        ['Descent 150 m → 25 m',      '7:30 – 10:30', '~3.0 min', 'Open (anticipation)',
         '40 m/min descent. LED canvas grows in resolution below the gondola.'],
        ['Main show at 25 m',         '10:30 – 18:00', '~7.5 min', 'CLOSED both ways',
         'Gondola holds at 25 m. Full 55,575-node themed show. No hallway movement.'],
        ['Final descent to ground',   '18:00 – 18:40', '~0.7 min', 'Closed',
         'Show closes as the gondola lands at the platform.'],
        ['Photo dwell on platform',   '18:40 – 20:00', '~1.3 min', 'Open (exit / next entry)',
         'Guests stay on the platform for photos. Hallway reopens for the next batch entry; previous batch exits during the next batch\'s boarding window.'],
        ['LED-field dwell (optional)','20:00 – 30:00', 'up to 10 min', 'Field (open)',
         'After clearing the platform, guests may linger inside the LED ring itself for photos. The next batch is already boarding the gondola — the dwell is parallel to gondola throughput.'],
    ],
    col_widths=[3.5, 2.5, 1.8, 3.0, 5.2])
add_para(doc, '')
h3(doc, 'Why 20 minutes (and not less)')
bullet(doc, 'Ascent + descent at the documented balloon speeds (50 m/min up, 40 m/min down) consumes ~6 minutes round-trip at 150 m. Below that we are not selling a flight — we are selling a lift.')
bullet(doc, 'A 7.5-minute main show is the floor for a recognisable themed narrative. Shorter than that and the All-Access pass (three shows in a night) loses meaning.')
bullet(doc, 'Photo dwell + boarding turnover need ~3 minutes combined for a safely-managed gondola hand-off at 25 guests/batch.')
h3(doc, 'Why the hallway must close during the show')
add_para(doc,
    'The main show runs with the gondola held at 25 m, looking straight down '
    'at the 60 m LED ring. Any pedestrian movement on the LED hallway — even '
    'a single person — produces a moving dark shape across the canvas that '
    'the gondola sees clearly from 25 m above. Closing the hallway during '
    'the show is non-negotiable for show quality; the on-platform photo '
    'dwell is what makes that closure operationally invisible.')

# ── 6. CAPACITY IMPACT ───────────────────────────────────────────────
h1(doc, '6. Capacity impact')
add_para(doc,
    'The 25-guest cap reduces per-batch loading, but the compressed day '
    'and night cycles (16 min and 20 min respectively at 150 m) more '
    'than offset the cap. Both day and night seat capacities increase '
    'relative to the blueprint baseline.')
add_table(doc,
    header=['Capacity line', 'Baseline', 'Revised', 'Change'],
    rows=[
        ['Day seat capacity',   '6 × 30 = 180 / day',     '11 × 25 = 275 / day',    '+ 52.8 %'],
        ['Night seat capacity', '8 × 30 = 240 / night',   '16 × 25 = 400 / night',  '+ 66.7 %'],
        ['Annual seat days (250 vs 300)', '300 / year',   '250 / year',             '– 16.7 %'],
    ],
    col_widths=[5.0, 4.0, 4.0, 3.0])
add_para(doc,
    'Net: both day-side and night-side annual paid seats increase '
    'relative to baseline once the 250-day operating year is applied. '
    'Night ARPU remains >2× day ARPU, so the revenue mix still tilts '
    'toward night — but the day product now contributes meaningfully '
    'rather than being a token line item.',
    size=10, color=MUTED, italic=True)

# ── 7. PRICING ASSUMPTION ────────────────────────────────────────────
h1(doc, '7. Pricing assumption — ARPU held at baseline')
add_para(doc,
    'Day blended ARPU is held at AED 170 and night blended ARPU at AED 361, '
    'i.e. unchanged from the pitch-blueprint.md baseline. The defensible '
    'rationale at a 150 m peak is:')
bullet(doc, 'Night value is driven by the LED canvas show, the three museum-tied themes, and the Saadiyat night-destination ecosystem — not by raw altitude. The show is paid for; the altitude is the delivery vehicle.')
bullet(doc, '150 m sits within the global peer band for tethered/observation balloons: London Eye 135 m, Singapore Flyer 165 m. The premium-experience price point is defensible at this band.')
bullet(doc, 'The All-Access pass (AED 695) is grounded in seeing all three themed shows in one evening. Altitude is irrelevant to the upsell.')
bullet(doc, 'Day flight pricing (AED 195 adult) is justified by the unobstructed museum-and-skyline panorama at 150 m, not by altitude alone.')
add_table(doc,
    header=['ARPU line', 'Baseline 300 m', 'Revised 150 m', 'Change'],
    rows=[
        ['Day blended ARPU',   'AED 170', 'AED 170', '0 %'],
        ['Night blended ARPU', 'AED 361', 'AED 361', '0 %'],
    ],
    col_widths=[5.5, 3.5, 3.5, 3.5])
add_para(doc,
    'A moderate ARPU haircut (AED 140 day / AED 320 night) is carried only '
    'as a downside sensitivity in section 9 for prudence — it is not the '
    'headline model.',
    size=10, color=MUTED, italic=True)

# ── 8. REVENUE BUILD ─────────────────────────────────────────────────
h1(doc, '8. Revenue build — formulas and worked numbers')
h3(doc, 'Ticket revenue formulas')
add_para(doc, 'Day ticket revenue =', bold=True)
add_para(doc, '   11 rides/day × 25 guests × load × 250 operating days × AED 170 ARPU',
         italic=True, color=MUTED)
add_para(doc, 'Night ticket revenue =', bold=True)
add_para(doc, '   16 batches/night × 25 guests × load × 250 operating nights × AED 361 ARPU ÷ 1.4 rides per unique guest',
         italic=True, color=MUTED)
h3(doc, 'Worked example — Base Case (58 % load)')
bullet(doc, 'Day:  11 × 25 × 0.58 × 250 × 170 = AED 6.78 M  ≈  AED 6.8 M')
bullet(doc, 'Night: (16 × 25 × 0.58 × 250) ÷ 1.4 × 361 = (58,000 ÷ 1.4) × 361 = 41,429 unique guests × AED 361 = AED 14.96 M  ≈  AED 15.0 M')
bullet(doc, 'F&B uplift from 10-min LED-field dwell: ~+30 % on baseline F&B per scenario.')
h3(doc, 'Load factors used per scenario')
add_table(doc,
    header=['Scenario', 'Load', 'Positioning'],
    rows=[
        ['Conservative', '32 %', 'Ramp-up / weather-constrained resilience case.'],
        ['Base Case',    '58 %', 'Practical steady operating case below the baseline 70 % benchmark.'],
        ['Strong',       '73 %', 'Demand-strong case, still below the baseline full-upside forecast.'],
    ],
    col_widths=[3.5, 2.0, 10.5])
h3(doc, 'Revenue by scenario (annual)')
add_table(doc,
    header=['Revenue stream', 'Conservative', 'Base Case', 'Strong'],
    rows=[
        ['Day ticket revenue (16-min cycle)',  'AED 3.7 M',  'AED 6.8 M',  'AED 8.5 M'],
        ['Night ticket revenue',               'AED 8.3 M',  'AED 15.0 M', 'AED 18.8 M'],
        ['F&B / retail (with LED-field dwell)','AED 0.4 M',  'AED 0.9 M',  'AED 1.4 M'],
        ['Sponsorship — content + branding',   'AED 0.7 M',  'AED 1.6 M',  'AED 2.2 M'],
        ['Museum partnership + cross-sell',    'AED 0.1 M',  'AED 0.4 M',  'AED 0.5 M'],
        ['Private gondola + VIP events',       'AED 0.4 M',  'AED 0.8 M',  'AED 1.0 M'],
        ['Total annual revenue',               'AED 13.6 M', 'AED 25.5 M', 'AED 32.4 M'],
    ],
    col_widths=[5.5, 3.0, 3.0, 3.0])
add_para(doc,
    'Day ticket revenue nearly doubles vs. the previous version of this '
    'memo (was AED 3.7 M base; now AED 6.8 M) on the back of the 16-min '
    'cycle. F&B is uplifted by ~30 % from the LED-field dwell. All other '
    'non-ticket lines are held flat — they remain a conservative ramp.',
    size=10, color=MUTED, italic=True)

# ── 9. EBITDA & PAYBACK ──────────────────────────────────────────────
h1(doc, '9. EBITDA and payback')
add_para(doc,
    'OpEx is held close to the baseline structure documented in '
    'pitch-blueprint.md Part 4, with a small uplift for the additional '
    'day operating hours (more crew rotation) — AED 9.6 M / 9.9 M / '
    '10.3 M across Conservative / Base / Strong.')
add_table(doc,
    header=['Scenario', 'Revenue', 'OpEx', 'EBITDA', 'Margin', 'Simple payback'],
    rows=[
        ['Conservative', 'AED 13.6 M', 'AED 9.6 M',  'AED 4.0 M',  '29 %', '~6.5 yrs'],
        ['Base Case',    'AED 25.5 M', 'AED 9.9 M',  'AED 15.6 M', '61 %', '~1.7 yrs'],
        ['Strong',       'AED 32.4 M', 'AED 10.3 M', 'AED 22.1 M', '68 %', '~1.2 yrs'],
    ],
    col_widths=[3.0, 2.5, 2.5, 2.5, 2.0, 3.5])
add_para(doc, '')
add_para(doc,
    'Simple payback is computed against the AED 25.9 M aluminum-tier CAPEX. '
    'Below the conservative case the model loses payback discipline very '
    'quickly — see section 11 for the night-window sensitivity.',
    size=10, color=MUTED, italic=True)

# ── 10. 10-YEAR RETURN ESTIMATE ──────────────────────────────────────
h1(doc, '10. 10-year return estimate (Base Case)')
add_para(doc, 'The 10-year cash-flow ramp uses these assumptions:')
bullet(doc, 'Year 1 lands at the Conservative ramp-up EBITDA.')
bullet(doc, 'Year 2 reaches 85 % of Base EBITDA.')
bullet(doc, 'Year 3 reaches Base EBITDA.')
bullet(doc, 'Years 4 – 10 grow at 4 % per year.')
bullet(doc, 'AED 2.5 M balloon envelope replacement reserves are booked in Years 5 and 10.')
add_table(doc,
    header=['Metric', 'Revised Base Case'],
    rows=[
        ['10-year IRR',          '~45 %'],
        ['10-year NPV at 8 %',   '~AED 73.4 M'],
        ['10-year MOIC',         '~6.0×'],
    ],
    col_widths=[8.0, 6.0])
add_para(doc, '')
h3(doc, '10-year cash flow ladder (Base Case, EBITDA in AED M)')
add_table(doc,
    header=['Year', 'EBITDA', 'Balloon CapEx', 'Net cash flow'],
    rows=[
        ['1',  '4.0',   '—',       '4.0'],
        ['2',  '13.3',  '—',       '13.3'],
        ['3',  '15.6',  '—',       '15.6'],
        ['4',  '16.22', '—',       '16.22'],
        ['5',  '16.87', '(2.5)',   '14.37'],
        ['6',  '17.55', '—',       '17.55'],
        ['7',  '18.25', '—',       '18.25'],
        ['8',  '18.98', '—',       '18.98'],
        ['9',  '19.74', '—',       '19.74'],
        ['10', '20.53', '(2.5)',   '18.03'],
        ['10-yr total', '', '',    'AED 156.0 M'],
    ],
    col_widths=[2.5, 3.0, 3.5, 5.0])

# ── 11. DOWNSIDE SENSITIVITY — REDUCED ARPU ──────────────────────────
h1(doc, '11. Downside sensitivity — reduced 150 m ARPU')
add_para(doc,
    'If the executive review wants a more conservative pricing posture, '
    'the model below applies a moderate ARPU haircut. This is carried for '
    'prudence; it is not the recommended headline assumption.')
add_table(doc,
    header=['ARPU line', 'Main case', 'Downside sensitivity'],
    rows=[
        ['Day blended ARPU',   'AED 170', 'AED 140'],
        ['Night blended ARPU', 'AED 361', 'AED 320'],
    ],
    col_widths=[6.0, 4.0, 4.0])
add_table(doc,
    header=['Scenario', 'Revenue', 'EBITDA', 'Margin', 'Simple payback'],
    rows=[
        ['Conservative', 'AED 12.0 M', 'AED 2.4 M',  '20 %', '~10.8 yrs'],
        ['Base Case',    'AED 22.6 M', 'AED 12.7 M', '56 %', '~2.0 yrs'],
        ['Strong',       'AED 28.8 M', 'AED 18.5 M', '64 %', '~1.4 yrs'],
    ],
    col_widths=[3.5, 3.0, 3.0, 2.5, 3.5])
add_para(doc,
    'Base-case 10-year returns under this downside pricing sensitivity are '
    'approximately 38 % IRR, AED 56 M NPV at 8 %, and 4.9× MOIC. The '
    'downside case is now meaningfully more resilient than in the prior '
    'version of this memo because day-side revenue is larger and less '
    'sensitive to ARPU (the day product trades on the panorama, not the '
    'show).',
    italic=True, color=MUTED, size=10)

# ── 12. NIGHT WINDOW SENSITIVITY ─────────────────────────────────────
h1(doc, '12. Sensitivity — night operating window')
add_para(doc,
    'The 20-minute cycle is the central operational lever of the revised '
    'model. If it is approved but the nightly operating window is shorter '
    'than the assumed 5.5 hours, the headline returns move materially.')
add_table(doc,
    header=['Night batches / night', 'Base revenue', 'Base EBITDA', 'Simple payback', '10-year IRR'],
    rows=[
        ['9 batches',  'AED 18.9 M', 'AED 9.0 M',  '~2.9 yrs', '~22 %'],
        ['12 batches', 'AED 21.7 M', 'AED 11.9 M', '~2.2 yrs', '~32 %'],
        ['15 batches', 'AED 24.6 M', 'AED 14.7 M', '~1.8 yrs', '~42 %'],
        ['16 batches', 'AED 25.5 M', 'AED 15.6 M', '~1.7 yrs', '~45 %'],
    ],
    col_widths=[3.5, 3.0, 3.0, 3.0, 3.0])
add_para(doc,
    'The model assumes 16 batches/night. The 16-min day cycle materially '
    'cushions this sensitivity vs. the prior version of this memo — even '
    'at 9 night batches the project clears a ~22 % 10-year IRR because '
    'day-side revenue is now contributing ~AED 6.8 M of Base Case income.',
    size=10, color=MUTED, italic=True)

# ── 13. KEY TAKEAWAYS ────────────────────────────────────────────────
h1(doc, '13. Key takeaways for the executive read')
bullet(doc, '150 m is the true new peak. The deck and this memo are consistent on this point — the attraction is no longer marketed as 300 m capable.')
bullet(doc, 'The 25-guest cap improves guest comfort and gives operating margin. Same gondola, same certification — the loading uplift vs. Dubai actuals (12 pax) is positioned as a demand-side claim from the LED canvas product, not a hardware re-rating.')
bullet(doc, 'The 16-minute day cycle is enabled by the 150 m peak shortening vertical motion to ~6.75 min round-trip. Day capacity goes from baseline 180 / day to 275 / day. Peak panoramic hold retained at 5 min so the photo moment is not compromised.')
bullet(doc, 'The 20-minute night-batch cadence more than doubles nightly batch count vs. the blueprint (8 → 16 batches).')
bullet(doc, 'After the gondola lands, guests may linger inside the LED ring for up to 10 minutes for photos. The next batch is already boarding the gondola — the dwell does not reduce throughput. Lifts F&B opportunity and strengthens the All-Access pass narrative.')
bullet(doc, 'The LED hallway closes during the main show. The on-platform photo dwell + LED-field dwell on the exit side make that closure operationally invisible.')
bullet(doc, 'ARPU is held at the baseline AED 170 day / AED 361 night because value is driven by the LED show + museum themes + Saadiyat ecosystem, not by raw altitude. A moderate ARPU haircut is carried as a downside sensitivity only.')
bullet(doc, 'The revised base case strengthens the investment story: AED 25.5 M revenue, AED 15.6 M EBITDA (61 % margin), ~1.7-year payback, ~45 % 10-year IRR, AED 73.4 M NPV at 8 %, 6.0× MOIC.')

# ── 14. APPENDIX A — DECK ALIGNMENT ──────────────────────────────────
h1(doc, 'Appendix A — Cross-reference to PitchDeck v2 (index.html)')
add_para(doc,
    'This appendix lists every headline figure in the revised deck that '
    'this memo is the source of. Where the deck currently shows a value '
    'inherited from an earlier revision and not yet updated to the figures '
    'below, that gap is flagged for follow-up.')
add_table(doc,
    header=['Deck location', 'Headline figure (this memo)', 'Notes'],
    rows=[
        ['Slide 1 hero — altitude stat',           '150 m',                 'Replaces previous 300 m stat.'],
        ['Slide 3 — night ride cycle length',      '20-minute batch',       'Replaces previous ~41-minute single cycle.'],
        ['Slide 3 — peak altitude step',           '150 m',                 'Replaces previous 300 m peak.'],
        ['Slide 3 — descent narrative',            '125 m at 40 m/min',     'Replaces previous "275 m at 40 m/min".'],
        ['Slide 3 — finale narrative',             'Photo dwell on platform',
                                                  'New beat: hallway-closed-during-show is operationally invisible because guests linger on the platform.'],
        ['Slide 6 — CAPEX (single-spec aluminum)', 'AED 25.9 M',            'Unchanged from blueprint.'],
        ['Slide 7 — annual revenue (Base)',        'AED 25.5 M',            'Updated for 16-min day cycle and LED-field dwell.'],
        ['Slide 7 — EBITDA (Base)',                'AED 15.6 M',            'Updated.'],
        ['Slide 7 — EBITDA margin (Base)',         '61 %',                  'Updated.'],
        ['Slide 7 — simple payback (Base)',        '~1.7 yrs',              'Updated.'],
        ['Slide 7 — 10-year IRR (Base)',           '~45 %',                 'Updated.'],
        ['Slide 7 — 10-year NPV at 8 %',           'AED 73.4 M',            'Updated.'],
        ['Slide 7 — 10-year MOIC',                 '6.0×',                  'Updated.'],
    ],
    col_widths=[5.0, 4.5, 6.5])

# Footer line
add_para(doc, '')
add_para(doc,
    '— end of memo —',
    italic=True, color=MUTED, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

# Save
out = r'd:\Proposal-pitch\financial-model-revision-150m-v2.docx'
doc.save(out)
print('Wrote ' + out)
