"""
Generate Original-Proposal-vs-Current-Pitch-Comparison.docx

A complete side-by-side comparison of:
  - The original Hero Experiences technical proposal
    (archive/ADB_at_Saadiyat_Technical_Proposal.pdf)
  - The current Canvas From The Sky pitch
    (index.html / pitch-blueprint.md / financial-model-revision-150m-v2)

Includes the 12% Partner / 83% Operator / 5% VAT revenue-share waterfall
from the original (page 11) — the structural element the current pitch
deck does not yet incorporate.

Output: Original-vs-Current-Comparison.docx in the project root.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GOLD   = RGBColor(0x8C, 0x6F, 0x2A)
TEXT   = RGBColor(0x1A, 0x1A, 0x1A)
MUTED  = RGBColor(0x55, 0x55, 0x55)
RED    = RGBColor(0xB0, 0x2A, 0x2A)
GREEN  = RGBColor(0x2A, 0x6F, 0x3A)


def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_para(doc, text, *, bold=False, italic=False, size=11, color=TEXT,
             align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = TEXT
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = GOLD
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = TEXT
    return p


def add_table(doc, header, rows, *, col_widths=None, first_col_left=True,
              section_marker='__SECTION__'):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    # Header
    for i, txt in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, 'F1E9D2')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(txt))
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = TEXT
    # Body
    for r, row in enumerate(rows, start=1):
        first = str(row[0])
        if first.startswith(section_marker):
            label = first[len(section_marker):].strip()
            cells = table.rows[r].cells
            merged = cells[0]
            for c in cells[1:]:
                merged = merged.merge(c)
            set_cell_shading(merged, 'EFE2C2')
            merged.text = ''
            p = merged.paragraphs[0]
            run = p.add_run(label)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT
            continue
        for i, val in enumerate(row):
            cell = table.rows[r].cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT
            if i == 0 and first_col_left:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif i > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return table


def bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT
    return p


def callout(doc, title, body, *, color=RED):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('⚠ ' + title)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = color
    add_para(doc, body, size=10, color=TEXT, space_after=6)


# ──────────────────────────────────────────────────────────────────────
# Build the document
# ──────────────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── COVER ────────────────────────────────────────────────────────────
add_para(doc, 'COMPARATIVE ANALYSIS',
         bold=True, size=10, color=MUTED)
add_para(doc, 'Original Proposal vs. Current Pitch',
         bold=True, size=24, color=TEXT)
add_para(doc, 'Abu Dhabi Balloon at Saadiyat  →  Canvas From The Sky',
         italic=True, size=12, color=MUTED)
add_para(doc, '')
add_para(doc,
    'Side-by-side comparison of the original Hero Experiences Group '
    'technical proposal and the current Canvas From The Sky pitch — '
    'including the 12% Partner / 83% Operator / 5% VAT revenue-share '
    'waterfall from the original (page 11) that the current pitch does '
    'not yet apply.',
    size=11, color=TEXT)
add_para(doc, '')

meta = doc.add_table(rows=4, cols=2)
meta.autofit = False
for i, w in enumerate([3.5, 12.5]):
    for cell in meta.columns[i].cells:
        cell.width = Cm(w)
meta_rows = [
    ('Date',     '18 May 2026'),
    ('Audience', 'CEO / executive review — internal'),
    ('Sources (original)', 'archive/ADB_at_Saadiyat_Technical_Proposal.pdf (19-page Hero Experiences proposal)'),
    ('Sources (current)',  'pitch-blueprint.md · financial-model-revision-150m-v2 · index.html'),
]
for i, (k, v) in enumerate(meta_rows):
    c0, c1 = meta.rows[i].cells
    c0.text = ''
    run = c0.paragraphs[0].add_run(k)
    run.bold = True; run.font.name = 'Calibri'; run.font.size = Pt(10); run.font.color.rgb = MUTED
    c1.text = ''
    run = c1.paragraphs[0].add_run(v)
    run.font.name = 'Calibri'; run.font.size = Pt(10); run.font.color.rgb = TEXT

# ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────────
h1(doc, '1. Executive summary')
add_para(doc,
    'The current Canvas From The Sky pitch is a rebuild of the original '
    'Hero Experiences proposal, not an update. Both pitches use the same '
    'balloon hardware, but everything else has changed: the product '
    'definition, the operating model, the ticket structure, the CAPEX '
    'scope, and the revenue lines.')

add_para(doc,
    'Most differences are deliberate and improve the model. One element '
    'from the original is missing from the current pitch and needs to be '
    'addressed before the next investor read:')

callout(doc,
    'The original revenue-share waterfall is not in the current model.',
    'Page 11 of the original proposal defines a 5% VAT / 12% Partner-or-'
    'Authority / 83% Operator structure that leaves the investor with the '
    'residual 17% of the post-VAT, post-Partner pool — roughly 14.2% of '
    'gross revenue. Applied to the current Base Case, this changes the '
    'investor return profile substantially (see §10).')

# ── 2. AT A GLANCE ───────────────────────────────────────────────────
h1(doc, '2. At a glance')
add_para(doc,
    'A single-table snapshot of every dimension that moved between the '
    'original proposal and the current pitch. Detailed comparisons follow '
    'in §3 onward.')

add_table(doc,
    header=['Dimension', 'Original (Hero proposal)', 'Current (Canvas From The Sky)', 'Direction of change'],
    rows=[
        ['__SECTION__Concept & product'],
        ['Project name',                  'Abu Dhabi Balloon at Saadiyat',     'Canvas From The Sky',            'Rebranded'],
        ['Core product',                  'Tethered helium balloon ride',      'Balloon ride + LED canvas show', 'Added LED night layer'],
        ['Day product',                   'Sightseeing 10 hrs',                'Sightseeing 9am–12pm (3 hrs)',   'Compressed to mornings'],
        ['Night product',                 'None',                              '3 themed LED shows',              'New product line'],
        ['Museum partnerships',           'Mentioned (combo tickets only)',    'Programmed (cross-discount + Cultural Pass)', 'Formalised'],
        ['__SECTION__Hardware (unchanged)'],
        ['Balloon system',                'Existing Dubai Balloon (relocated)', 'Same balloon hardware',          'Unchanged'],
        ['Envelope volume / lift',        '6,200 m³ / 4,800 kg',                'Same',                           'Unchanged'],
        ['Gondola spec (envelope)',       '30 pax + 2 pilots',                  'Same',                           'Unchanged'],
        ['Tether max length',             '300 m',                              '300 m (capable); 150 m marketed', 'Marketed altitude lowered'],
        ['__SECTION__Operating model'],
        ['Peak operating altitude',       '300 m',                              '150 m',                          'Reduced — true new peak'],
        ['Commercial pax / flight',       '12 (page 9 model)',                  '25 (v2 memo) / 30 (blueprint)',  'Increased loading'],
        ['Day cycle',                     '10 min (5 flights/hr)',              '30 min (2 flights/hr)',          'Honest at 300 m; was unrealistic'],
        ['Day flights/day',               '50 (10 hrs × 5/hr)',                 '6 (3-hr window)',                'Window compressed for night product'],
        ['Night cycle',                   'N/A',                                '20 min batches (3/hr × 5.5 hrs)', 'New cadence'],
        ['Operating days',                '365',                                '250 (v2) / 300 (blueprint)',     'More conservative'],
        ['Weather cancellation',          '15%',                                '~32% (v2) / 15% (blueprint)',    'More conservative'],
        ['Load factor (base)',            '70%',                                '58% (v2) / 70% (blueprint)',     'More conservative'],
        ['Effective day guests/yr',       '~128,520',                            '~21,750 (v2) / 37,800 (blueprint)', 'Lower (day-window cut)'],
        ['Effective night guests/yr',     'N/A',                                 '~58,000 (v2) / 36,000 (blueprint)', 'New stream'],
        ['__SECTION__Tickets'],
        ['Day Adult',                     'AED 195',                             'AED 195',                        'Unchanged'],
        ['Day Child',                     'AED 125',                             'AED 125',                        'Unchanged'],
        ['Express Pass',                  'AED 295 (10% mix)',                   'Discontinued',                    'Removed'],
        ['Night Adult',                   'N/A',                                 'AED 295',                         'New'],
        ['Night Child',                   'N/A',                                 'AED 175',                         'New'],
        ['All-Access Pass',               'N/A',                                 'AED 695 (3 shows)',               'New'],
        ['Private Gondola',               'Mentioned, no price',                 'AED 20,000',                       'Priced'],
        ['Day blended ARPU',              'AED 192',                             'AED 170',                          '–AED 22 (more children)'],
        ['Night blended ARPU',            'N/A',                                 'AED 361',                          'New'],
        ['__SECTION__Financial outputs (Year 1, Base)'],
        ['Annual revenue',                'AED 24.7 M',                          'AED 22.2 M (v2) / 24.7 M (blueprint)', 'Similar headline'],
        ['Revenue streams',               '1 (day tickets only)',                '6 (day, night, F&B, sponsor, museum, VIP)', 'Diversified'],
        ['Annual OpEx',                   'AED 7.2 M',                           'AED 9.8 M',                        '+AED 2.6 M (LED layer)'],
        ['EBITDA',                        'Not stated (implied ~AED 17.5 M)',    'AED 12.4 M (v2) / 14.9 M (blueprint)', 'Lower at project level'],
        ['CAPEX',                         'AED 15 M (relocation + civils)',      'AED 25.9 M (balloon + LED + content)', '+AED 10.9 M'],
        ['__SECTION__Revenue-share waterfall'],
        ['VAT',                           '5% of gross',                          'Not modelled',                     'GAP'],
        ['Partner / Authority',           '12% of net-of-VAT',                    'Not modelled',                     'GAP'],
        ['Operator (Hero)',               '83% of net-of-VAT-and-partner',        'Not modelled',                     'GAP'],
        ['Investor retained share',       '~14.2% of gross (residual 17%)',       'Not modelled (project-level only)', 'GAP'],
        ['10-yr investor retained',       'AED 35.4 M',                            'See §10 — restated',               'Needs restatement'],
        ['__SECTION__Returns (project-level, no waterfall)'],
        ['Simple payback',                'Implied ~6 yrs at investor level',     '~2.1 yrs (v2 project level)',       'Restated higher'],
        ['10-year NPV @ 8%',              'Not stated',                            'AED 51.4 M (v2) / 85 M (blueprint)', 'Project-level only'],
        ['10-year MOIC',                  'Not stated',                            '~4.7× (v2) / 5.5× (blueprint)',     'Project-level only'],
    ],
    col_widths=[4.5, 4.5, 4.5, 3.0])

# ── 3. CONCEPT & POSITIONING ─────────────────────────────────────────
h1(doc, '3. Concept & positioning')

h3(doc, 'Original — Hero proposal')
add_para(doc,
    '"The Abu Dhabi Balloon at Saadiyat" is positioned as a relocation of '
    'the proven Dubai Balloon at Atlantis, The Palm. The pitch is built '
    'around the existing certified system being ready to redeploy, with '
    'Saadiyat Island offering "the next stage" for a proven attraction. '
    'The product is a single tethered helium balloon ride to ~300 m for '
    '360° views of Saadiyat\'s cultural district.')

h3(doc, 'Current — Canvas From The Sky')
add_para(doc,
    'The same balloon hardware is wrapped inside a larger destination '
    'concept. A programmable 60 m LED canvas ring is added around the '
    'ground platform, turning the attraction into a two-product offer: '
    'daytime sightseeing (9am–12pm) and premium night LED shows tied to '
    'the three Saadiyat museums (Louvre, Zayed NM, Guggenheim). The '
    'balloon is reframed as "the aerial connector of Saadiyat\'s cultural '
    'spine" rather than a standalone ride.')

h3(doc, 'What this means commercially')
add_para(doc,
    'The original sold one product to one audience (daytime tourist '
    'sightseeing). The current pitch sells two products to two audiences '
    '(day family/tourist + night premium/cultural), with the second '
    'audience commanding 2× the ARPU of the first. The LED canvas is the '
    'reason a second audience exists. This is the load-bearing change in '
    'the entire revision.')

# ── 4. EXPERIENCE & OPERATING MODEL ──────────────────────────────────
h1(doc, '4. Experience & operating model')

h3(doc, 'Original cycle math (page 9)')
add_para(doc,
    'The original proposal documents a 10-minute flight cycle allowing '
    '5 flights per hour over a 10-hour operating day, with a 12-guest '
    'commercial load.')
bullet(doc, '5 flights/hr × 10 hrs = 50 flights/day max')
bullet(doc, '50 × 12 pax = 600 guests/day max')
bullet(doc, '600 × 0.85 (weather) × 0.70 (load) = 357 guests/day')
bullet(doc, '357 × 365 days = 128,520 guests/year')

callout(doc,
    'The original\'s 10-minute cycle is mathematically impossible at 300 m.',
    'At the documented balloon speeds (50 m/min ascent, 40 m/min descent), '
    'vertical motion alone consumes 13.5 minutes round-trip at 300 m. '
    'A 10-minute cycle inclusive of boarding, peak hold, and deboarding '
    'cannot physically exist. The realistic cycle at 300 m is ~30 minutes '
    '(2 flights/hr), as the current pitch uses. The original\'s headline '
    'AED 24.7 M revenue therefore rests on a cycle assumption that is '
    'overstated by 2.5×.')

h3(doc, 'Current cycle math (v2 memo)')
add_para(doc,
    'The current pitch uses two cycle definitions, one per product:')
bullet(doc, 'Day: 30 min cycle × 6 rides over 9am–12pm = 6 day flights')
bullet(doc, 'Night: 20 min batch cycle × 16 batches over sundown–midnight (feasible because peak altitude is 150 m, not 300 m)')

add_para(doc,
    'The 20-minute night batch is only feasible because the marketed peak '
    'altitude was reduced from 300 m to 150 m. At 150 m, vertical motion '
    'is 6.75 minutes round-trip — enough headroom for a 7.5-minute main '
    'show at 25 m show altitude plus boarding turnover.')

h3(doc, 'Gondola loading — the silent shift')
add_para(doc,
    'The most consequential operating-model change is the commercial '
    'loading assumption. The original page-6 spec sheet says 30 passengers; '
    'the original page-9 revenue model assumes 12. The current pitch uses '
    '25 (v2 memo) or 30 (blueprint). This is a 2–2.5× increase in '
    'per-flight loading vs. the original\'s commercial assumption.')

add_table(doc,
    header=['Loading assumption', 'Per flight', 'Vs. envelope spec', 'Implied effective occupancy'],
    rows=[
        ['Original envelope spec',           '30',  '100%',  '—'],
        ['Original revenue model (page 9)',  '12',  '40%',   '12 × 0.70 = 8.4 pax/flight = 28% of envelope'],
        ['Current blueprint',                '30',  '100%',  '30 × 0.70 = 21 pax/flight = 70% of envelope'],
        ['Current v2 memo',                  '25',  '83%',   '25 × 0.58 = 14.5 pax/flight = 48% of envelope'],
    ],
    col_widths=[5.5, 2.5, 3.0, 5.0])

add_para(doc,
    'The likely rationale for the original\'s 12-guest assumption is '
    'real-world Dubai Balloon operating averages — Hero has 250,000+ '
    'safe flights of historical data. A 30-pax envelope averaging '
    '12 pax/flight is a ~40% effective load baked into the loading line. '
    'The current pitch raises this on the thesis that the LED canvas + '
    'night product + museum ecosystem will drive demand uplift. '
    'A defensible framing for investors is: same gondola, same safety '
    'envelope; loading uplift is a demand-side claim attributable to the '
    'LED canvas, not a hardware change.',
    italic=True, color=MUTED, size=10)

# ── 5. DAYTIME REVENUE — APPLES-TO-APPLES ────────────────────────────
h1(doc, '5. Daytime revenue — apples-to-apples')
add_para(doc,
    'If both models ran daytime sightseeing only on equal hours and equal '
    'realistic cycle assumptions, the revenue comparison is as follows.')

h3(doc, 'Scenario A — accept original\'s 10-min cycle (50 flights/day)')
add_table(doc,
    header=['Model', 'Flights/day', 'Pax/flight', 'Effective load', 'ARPU', 'Op days', 'Annual revenue'],
    rows=[
        ['Original (page 9)',   '50',  '12',  '0.85 × 0.70',  'AED 192',  '365',  'AED 25.1 M'],
        ['Current, same cadence', '50', '25', '0.58',         'AED 170',  '250',  'AED 30.8 M'],
    ],
    col_widths=[4.0, 2.0, 2.0, 2.5, 2.0, 1.5, 2.5])

h3(doc, 'Scenario B — realistic 30-min cycle at 300 m (20 flights/day)')
add_table(doc,
    header=['Model', 'Flights/day', 'Pax/flight', 'Effective load', 'ARPU', 'Op days', 'Annual revenue'],
    rows=[
        ['Original config (honest cycle)',     '20',  '12',  '0.85 × 0.70',  'AED 192',  '365',  'AED 10.0 M'],
        ['Current config, full 10-hr day',     '20',  '25',  '0.58',         'AED 170',  '250',  'AED 12.3 M'],
        ['Current config, actual (3-hr day)',  '6',   '25',  '0.58',         'AED 170',  '250',  'AED 3.7 M'],
    ],
    col_widths=[5.0, 2.0, 2.0, 2.5, 2.0, 1.5, 2.5])

h3(doc, 'Read')
bullet(doc,
    'Per realistic flight, the current pitch produces ~+23% revenue '
    'vs. the original at equal operating hours, driven by the larger '
    'commercial loading (25 vs 12). ARPU is slightly lower (AED 170 vs '
    '192) because the current day mix has 35% children vs. the original\'s '
    '20%.')
bullet(doc,
    'The current pitch runs only 3 daytime hours instead of 10, which is '
    'why daytime revenue is AED 3.7 M and not AED 12.3 M. The afternoon '
    'window is intentionally kept dark to protect the premium night '
    'product. This is a deliberate trade-off, not a hardware constraint — '
    'and represents ~AED 8.6 M of unbooked daytime upside if the trade '
    'were reversed.')

# ── 6. CAPEX ─────────────────────────────────────────────────────────
h1(doc, '6. CAPEX comparison')

h3(doc, 'Original CAPEX scope (AED 15 M, page 11)')
bullet(doc, 'Relocation, transport, and reassembly of existing tethered helium balloon system')
bullet(doc, 'Site preparation, foundations, and utilities (power and minor infrastructure)')
bullet(doc, 'Basic guest facilities: ticketing kiosk, shaded waiting area, minimal landscaping')
bullet(doc, 'Working capital for launch marketing and pre-opening readiness')

h3(doc, 'Current CAPEX scope (AED 25.9 M, aluminum tier)')
add_table(doc,
    header=['Line', 'Amount', 'Notes'],
    rows=[
        ['Balloon (gondola + envelope + winch system)',          'AED 8.0 M',   'Hardware only'],
        ['LED canvas (8 mm aluminum tile + adjustable feet)',    'AED 15.1 M',  'New layer; absorbs site civils'],
        ['Content launch package (3 hero shows + ambient)',      'AED 2.8 M',   'New layer'],
        ['Total',                                                 'AED 25.9 M', ''],
    ],
    col_widths=[7.5, 3.0, 6.5])

callout(doc,
    'The AED 8 M balloon line is undocumented.',
    'The current pitch books the balloon at AED 8 M but the source of '
    'that figure is not recorded in pitch-blueprint.md or any supporting '
    'memo. The most internally-consistent reading is "balloon hardware '
    'only; site civils are absorbed inside the AED 15.1 M LED canvas line '
    '(the canvas IS the ground installation)." If that\'s correct, the '
    'AED 25.9 M project CAPEX is genuinely comparable to "original AED 15 '
    'M relocation + AED 15 M LED canvas," modulo a small overlap. '
    'Recommendation: confirm AED 8 M with Hero (new procurement vs. '
    'relocation quote) and add a footnote to the deck disclosing scope.')

# ── 7. OPEX ──────────────────────────────────────────────────────────
h1(doc, '7. OpEx comparison')

add_table(doc,
    header=['Line', 'Original', 'Current', 'Notes'],
    rows=[
        ['Balloon operations (helium, crew, certs)', 'AED 7.2 M (combined)',  'AED 4.3 M',  'Lower in current — original bundled marketing/admin'],
        ['LED field maintenance',                    '—',                     'AED 0.4 M',  'New layer'],
        ['Content operations (net)',                 '—',                     'AED 2.6 M',  'New layer'],
        ['Site staff',                               '(in AED 7.2 M)',        'AED 1.2 M',  'Broken out'],
        ['Marketing & PR',                           '(in AED 7.2 M)',        'AED 0.5 M',  'Broken out'],
        ['Insurance / regulatory / admin',           '(in AED 7.2 M)',        'AED 0.5 M',  'Broken out'],
        ['LED power',                                '—',                     'AED 0.025 M', 'New layer'],
        ['Contingency (5%)',                         '—',                     'AED 0.25 M', 'Best-practice add'],
        ['Total annual OpEx',                        'AED 7.2 M',             'AED 9.8 M',  '+AED 2.6 M for the LED layer'],
    ],
    col_widths=[5.5, 3.0, 3.0, 5.0])

add_para(doc,
    'The +AED 2.6 M delta is the price of running the night LED product. '
    'In exchange the current pitch adds AED 15 M+ in night-ticket and '
    'sponsorship revenue. The LED layer is OpEx-positive by a wide margin '
    'at the project level — the question is whether that holds at the '
    'investor level after the revenue-share waterfall is applied (§10).',
    italic=True, color=MUTED, size=10)

# ── 8. REVENUE BUILD ─────────────────────────────────────────────────
h1(doc, '8. Revenue build — line by line')

h3(doc, 'Original (page 9) — single stream')
add_para(doc,
    'The original proposal has one revenue line: day ticket sales.')
add_table(doc,
    header=['Stream', 'Calculation', 'Year 1 (AED)'],
    rows=[
        ['Day ticket revenue',
         '128,520 guests × AED 192 weighted ARPU',
         'AED 24.7 M'],
        ['Total annual revenue', '', 'AED 24.7 M'],
    ],
    col_widths=[5.0, 7.0, 4.0])

h3(doc, 'Current (v2 memo) — six streams')
add_table(doc,
    header=['Stream', 'Conservative', 'Base Case', 'Strong'],
    rows=[
        ['Day ticket revenue',     'AED 2.0 M',  'AED 3.7 M',   'AED 4.7 M'],
        ['Night ticket revenue',   'AED 8.3 M',  'AED 15.0 M',  'AED 18.8 M'],
        ['F&B / retail',           'AED 0.3 M',  'AED 0.7 M',   'AED 1.1 M'],
        ['Sponsorship',            'AED 0.7 M',  'AED 1.6 M',   'AED 2.2 M'],
        ['Museum partnership',     'AED 0.1 M',  'AED 0.4 M',   'AED 0.5 M'],
        ['Private gondola / VIP',  'AED 0.4 M',  'AED 0.8 M',   'AED 1.0 M'],
        ['Total annual revenue',   'AED 11.8 M', 'AED 22.2 M',  'AED 28.3 M'],
    ],
    col_widths=[5.0, 3.5, 3.5, 4.0])

add_para(doc,
    'The current pitch headline (AED 22.2 M v2 / AED 24.7 M blueprint) '
    'is similar to the original (AED 24.7 M), but the composition is '
    'completely different: the original was 100% sightseeing tickets, '
    'while the current model is ~17% day tickets / ~68% night tickets / '
    '~15% ancillaries.',
    italic=True, color=MUTED, size=10)

# ── 9. THE REVENUE-SHARE WATERFALL (PAGE 11) ─────────────────────────
h1(doc, '9. Revenue-share waterfall — the missing structural piece')

add_para(doc,
    'Page 11 of the original proposal defines a revenue-share structure '
    'between four parties. The current pitch does not yet incorporate '
    'this. Re-stated explicitly:')

h3(doc, 'Original waterfall — sequential')
add_table(doc,
    header=['Step', 'Party', 'Rule', 'Cumulative share of GROSS'],
    rows=[
        ['1', 'Government',  '5% VAT on gross',                                  '5.00%'],
        ['2', 'Partner / Authority', '12% of net-of-VAT (i.e. 12% × 95%)',      '11.40%'],
        ['3', 'Operator (Hero Experiences)', '83% of the post-VAT, post-Partner pool (i.e. 83% × 88% × 95%)', '69.40%'],
        ['4', 'Investor (residual)', '17% of the post-VAT, post-Partner pool (i.e. 17% × 88% × 95%)', '14.21%'],
        ['',  'Total',        '',                                                 '100.00%'],
    ],
    col_widths=[1.5, 5.0, 7.0, 3.5])

add_para(doc,
    'Sanity check against the original\'s stated "AED 35.4 M investor '
    'retained over 10 years":')
bullet(doc, 'Year 1 gross revenue: AED 24.7 M')
bullet(doc, 'Year 1 investor share: 24.7 × 14.21% = AED 3.51 M')
bullet(doc, '10 years × AED 3.51 M (flat) = AED 35.1 M ≈ AED 35.4 M ✓')

h3(doc, 'What the splits look like in cash')
add_table(doc,
    header=['Party', 'Year 1 cash (orig. AED 24.7 M)', 'Year 1 cash (current Base AED 22.2 M)'],
    rows=[
        ['Government (VAT)',          'AED 1.24 M',  'AED 1.11 M'],
        ['Partner / Authority',       'AED 2.81 M',  'AED 2.53 M'],
        ['Operator (Hero)',           'AED 17.14 M', 'AED 15.41 M'],
        ['Investor',                  'AED 3.51 M',  'AED 3.16 M'],
        ['Gross',                     'AED 24.70 M', 'AED 22.20 M'],
    ],
    col_widths=[5.0, 5.5, 5.5])

add_para(doc,
    'The Operator keeps AED 15–17 M annually before paying any OpEx out '
    'of its 69.4% share. The Investor receives AED 3.16–3.51 M annually '
    'and is also the party that funded CAPEX. These are not symmetric '
    'positions.',
    italic=True, color=MUTED, size=10)

# ── 10. WHAT THE WATERFALL DOES TO INVESTOR RETURNS ──────────────────
h1(doc, '10. Applying the waterfall to the current pitch')

add_para(doc,
    'Below is the current Base Case (v2 memo, AED 22.2 M annual revenue, '
    'AED 25.9 M CAPEX) restated with the original\'s waterfall applied. '
    'For completeness the original\'s own returns are restated the same '
    'way against its CAPEX of AED 15 M.')

add_table(doc,
    header=['Metric', 'Original (orig. CAPEX 15 M)', 'Current (CAPEX 25.9 M)', 'Delta'],
    rows=[
        ['Year 1 gross revenue',                          'AED 24.7 M',  'AED 22.2 M',  '–AED 2.5 M'],
        ['Year 1 investor cash (14.21% of gross)',        'AED 3.51 M',  'AED 3.16 M',  '–AED 0.35 M'],
        ['Simple payback (flat, no growth)',              '~4.3 yrs',    '~8.2 yrs',    '+3.9 yrs'],
        ['Simple payback (5% YoY growth)',                '~3.8 yrs',    '~6.9 yrs',    '+3.1 yrs'],
        ['10-yr investor cash (flat)',                    'AED 35.1 M',  'AED 31.6 M',  '–AED 3.5 M'],
        ['10-yr investor cash (5% YoY)',                  'AED 44.2 M',  'AED 39.7 M',  '–AED 4.5 M'],
        ['10-yr MOIC at investor level (5% YoY)',         '~2.9×',       '~1.5×',       '–1.4×'],
        ['vs. Current PROJECT-level MOIC (no waterfall)',  '—',          '~4.7×',       'Gap from waterfall'],
    ],
    col_widths=[6.5, 4.5, 4.5, 2.5])

callout(doc,
    'Headline finding.',
    'If the original waterfall (12% Partner / 83% Operator / 5% VAT) is '
    'preserved in the current deal, the investor sees AED ~3.16 M of cash '
    'per year against a AED 25.9 M CAPEX. That is roughly an 8-year flat '
    'payback and a ~1.5× 10-year MOIC at the investor level — versus the '
    '~2.1 year payback and ~4.7× MOIC the current deck shows at the '
    'project level. The two numbers describe different pots of money. '
    'The deck should be explicit about which one it is quoting.')

h3(doc, 'Implications')
bullet(doc,
    'If we keep the original waterfall, the current pitch overstates '
    'investor-facing returns by a factor of ~3× because it reports at the '
    'project level, not at the investor level.')
bullet(doc,
    'The investor in the current deal is also funding an additional '
    'AED 15 M of LED canvas + AED 2.8 M of content that did not exist in '
    'the original — but the waterfall does not adjust to compensate for '
    'that. This is the structural inconsistency to resolve before the '
    'next investor read.')
bullet(doc,
    'Three possible resolutions (not mutually exclusive): (a) renegotiate '
    'the waterfall to give the investor a larger residual reflecting the '
    'expanded CAPEX scope; (b) carve out the LED canvas revenue stream '
    '(night tickets, sponsorship) under a different waterfall than the '
    'balloon stream; (c) restructure as a two-investor model where the '
    'balloon investor and the LED canvas investor are different parties '
    'with different return profiles.')

# ── 11. WHAT TO DO NEXT ──────────────────────────────────────────────
h1(doc, '11. Open questions for the CEO')

add_table(doc,
    header=['#', 'Question', 'Why it matters'],
    rows=[
        ['1', 'Is the 12% Partner / 83% Operator / 5% VAT waterfall still binding for Canvas From The Sky, or is the partnership structure being renegotiated for the larger CAPEX scope?',
              'Determines whether the deck should quote project-level or investor-level returns.'],
        ['2', 'What is the actual basis for the AED 8 M balloon line (new procurement quote vs. relocation cost vs. placeholder)?',
              'Affects CAPEX credibility under investor diligence.'],
        ['3', 'Is the current 3-hour daytime window a hard constraint or a soft choice? Could we extend day operations to recover ~AED 8.6 M of unbooked revenue?',
              'Largest unforced upside in the current model.'],
        ['4', 'Should we explicitly position the 25-pax commercial cap as +20% loading uplift attributable to LED canvas demand (vs. the 12-pax Dubai actuals)?',
              'Reframes a quiet discrepancy as a demand-side claim investors can validate.'],
        ['5', 'Do we retain the original\'s 365-day operating year, or stay with v2\'s 250-day conservative envelope?',
              'Worth ~AED 5 M annual revenue at full waterfall.'],
    ],
    col_widths=[1.0, 7.5, 7.0])

# ── APPENDIX A — ORIGINAL PROPOSAL FULL TEXT EXTRACTS ────────────────
h1(doc, 'Appendix A — Original proposal: full content extract')

add_para(doc,
    'For reference, the complete content of the original 19-page Hero '
    'Experiences proposal (archive/ADB_at_Saadiyat_Technical_Proposal.pdf), '
    'organised by section.')

h2(doc, 'A.1  Executive Summary (page 3)')
add_para(doc,
    'The Helium Tethered Balloon Experience is a proven, premium attraction '
    'offering guests a serene 10-minute ascent up to approximately 300 m '
    'for an unforgettable 360° view. Successfully launched and operated as '
    'The Dubai Balloon at Atlantis, The Palm. Hero Experiences Group invites '
    'a co-investor to fund CAPEX and provide the site, while Hero delivers '
    'complete operational management. Returns are distributed through a '
    'long-term revenue-share agreement. The fully certified system is '
    'ready for relocation; the project can be activated within months.')

h2(doc, 'A.2  Experience Overview (page 5)')
add_para(doc,
    'Signature cultural attraction. ~10-minute ascent to ~300 m above sea '
    'level, 360° aerial view of Saadiyat Island. Uninterrupted panoramas '
    'of Louvre Abu Dhabi, Zayed National Museum, and the future Guggenheim '
    'Abu Dhabi. Premium, family-friendly, fully accessible.')
add_para(doc,
    '"Each flight cycle will last approximately ten minutes, including '
    'embarkation and disembarkation, allowing for five flights per hour '
    'throughout a ten-hour operating day. With a twelve-guest capacity '
    'per flight, the attraction will accommodate around 600 guests per '
    'day. Across a full operating year (average 30 days per month), this '
    'equates to approximately 216,000 guests annually, or roughly 219,000 '
    'guests on a 365-day schedule."',
    italic=True, color=MUTED, size=10)

h2(doc, 'A.3  Specifications (page 6)')
add_table(doc,
    header=['Component', 'Specification'],
    rows=[
        ['Total height of system',   '34 m (111 ft)'],
        ['Envelope diameter',        '22.46 m (72 ft)'],
        ['Envelope volume',          '6,200 m³ (220,000 ft³)'],
        ['Nominal lift',             '4,800 kg (10,560 lbs)'],
        ['Helium loss',              '<70 m³/month'],
        ['Useful envelope pressure', 'up to 900 Pa'],
        ['Crew',                     '2 pilots'],
        ['Passenger capacity',       '30 passengers'],
        ['Gondola material',         'composite + aluminum'],
        ['Gondola exterior diameter','5.80 m (19 ft)'],
        ['Gondola floor area',       '13 m² (140 ft²)'],
        ['Gondola weight',           '850 kg (1,870 lbs)'],
        ['Tether diameter',          '22 mm'],
        ['Tether resistance',        '45 tons (99,000 lbs)'],
        ['Tether useful length',     'up to 300 m (984 ft)'],
        ['Winch house dimensions',   '4 × 2.2 × 1.5 m'],
        ['Winch house weight',       '6 tons'],
        ['Winch power',              '45 kW'],
        ['Ascent speed',             '~50 m/min (164 ft/min)'],
        ['Descent speed',            '~40 m/min (131 ft/min)'],
    ],
    col_widths=[6.0, 10.0])

h2(doc, 'A.4  Highlights (page 7)')
bullet(doc, 'The premier 360° observatory in Abu Dhabi')
bullet(doc, 'Customisable flight times for private bookings — VIP guests, proposals, engagements, exclusive photography sessions')
bullet(doc, 'Sunrise, sunset, and night flights with ambient lighting and optional onboard music or narration')
bullet(doc, 'Supports sustainable, low-impact infrastructure (tethered balloon = smaller footprint)')
bullet(doc, 'Panoramic views of Saadiyat Island: Louvre, Zayed NM, future Guggenheim')
bullet(doc, 'Private charter options for special occasions and media shoots')
bullet(doc, 'Combo ticketing opportunities with Louvre Abu Dhabi, museums, resorts, Saadiyat hospitality partners')

h2(doc, 'A.5  Revenue Potential (page 9)')
add_table(doc,
    header=['Field', 'Value'],
    rows=[
        ['Experience type',         'Tethered helium hot air balloon'],
        ['Average load',            '12 guests per flight'],
        ['Flight frequency',        '5 flights/hour × 10 hours/day = 600 guests/day max'],
        ['Operating days',          '365 days per year'],
        ['Seasonality',             '6 high, 3 shoulder, 3 low months'],
        ['Weather impact',          '15% flight cancellations annually'],
        ['Load factor',             '70% capacity utilisation'],
        ['Effective throughput',    '600 × 0.85 × 0.70 = 357 guests/day (~128,520/year)'],
        ['Ticket — Adult (80%)',    'AED 195'],
        ['Ticket — Child (20%)',    'AED 125'],
        ['Ticket — Express (10%)',  'AED 295'],
        ['Weighted average',        '~AED 192.4 per guest'],
        ['Year 1 revenue',          'AED 24.7 M'],
        ['YoY growth',              '+5%'],
        ['Payback structure',       '10 years'],
    ],
    col_widths=[5.5, 10.5])

h2(doc, 'A.6  Investment Structure (page 10)')
add_para(doc,
    'Co-investment partnership between project investor (or landholding '
    'entity) and Hero Experiences Group.')
h3(doc, 'Investor finances CAPEX:')
bullet(doc, 'Relocation, transport, reassembly of the tethered helium balloon system')
bullet(doc, 'Site preparation, foundations, utilities (power and minor infrastructure)')
bullet(doc, 'Basic guest facilities (ticketing kiosk, shaded waiting area, minimal landscaping)')
bullet(doc, 'Working capital for launch marketing and pre-opening readiness')
bullet(doc, 'Land parcel within Saadiyat Island (proposed: 24°31′55″N 54°24′01″E)')
h3(doc, 'Hero Experiences Group delivers:')
bullet(doc, 'Full end-to-end operations (staffing, training, safety management, guest services, maintenance)')
bullet(doc, 'Ticketing, pricing, yield strategy (OTAs, hotel concierges, tour operators)')
bullet(doc, 'Marketing, branding, PR campaigns')
bullet(doc, 'Helium supply and safety certification (Civil Aviation compliance, insurance)')

h2(doc, 'A.7  Revenue Share Framework (page 11) ★ Critical')
add_para(doc,
    'Reproduced verbatim from the original proposal:')
add_table(doc,
    header=['Field', 'Value'],
    rows=[
        ['CAPEX (Investor, one-time)',     'AED 15,000,000 (balloon, inflation system, deck, tether, branding)'],
        ['OPEX (Hero input)',              'AED 600,000 / month  =  AED 7,200,000 / year'],
        ['VAT',                            '5% (Government levy)'],
        ['Partner / Authority',            '12% of net-of-VAT (shared operational revenue)'],
        ['Operator (Hero Experiences)',    '83% of gross, after VAT (5%) and Partner share'],
        ['Investor Retained (10-yr forecast, including 5% YoY growth)', 'AED 35.4 M'],
    ],
    col_widths=[6.0, 10.0])

h2(doc, 'A.8  Why Saadiyat, Why Now (pages 13–16)')
bullet(doc, 'Saadiyat is 2,700 ha master-planned cultural-and-leisure island; world-class institutions + luxury hospitality + beaches + conservation nature')
bullet(doc, '2024: 10% YoY visitor growth across hotels and museum venues; ~74% hotel occupancy; +14% ADR YoY')
bullet(doc, 'Aligned with Abu Dhabi DCT "Tourism Strategy 2030": 39.3 M annual visitors, AED 90 B national economy contribution')
bullet(doc, 'Multiple major museum openings (Zayed NM December 2025) — heightened visibility and global appeal')
bullet(doc, 'Demand rising not plateauing; cultural district in transformation; emergent not saturated — opportunity to own landmark status')
bullet(doc, 'Differentiation: balloon adds vertical vantage point; museums + resorts + beaches mostly ground-level')
bullet(doc, 'Stay-and-explore profile: arrive early for balloon → museum → sunset resort/beach → cross-sell ecosystem')
bullet(doc, 'First-mover brand advantage: position balloon as signature icon rather than late-entry amenity')

h2(doc, 'A.9  Why Partner With Hero Experiences Group (page 17)')
bullet(doc, 'GCC\'s leading operator of premium aerial and cultural attractions')
bullet(doc, 'Only company in the region with proven expertise across both helium tethered balloons and hot air balloon operations')
bullet(doc, 'Designed, launched, and operated the UAE\'s first permanent helium balloon attraction — 250,000+ safe passenger flights, flawless safety and compliance record')
bullet(doc, 'Hero Balloon Flights brand — one of the world\'s most recognised hot air balloon companies, GCAA-certified')
bullet(doc, 'Turn-key operational model: certification, staffing, ticketing, marketing, maintenance — all in-house')
bullet(doc, 'Decade of aviation expertise, award-winning brands, trusted reputation with government tourism authorities')

h2(doc, 'A.10  Closing — A Defining Opportunity for Abu Dhabi (page 18)')
add_para(doc,
    '"We are offering Abu Dhabi the first right of opportunity to secure a '
    'proven, world-class attraction... The system is fully built, certified, '
    'and ready to be relocated, requiring only land allocation and capital '
    'investment to bring it to life once more. Saadiyat Island represents '
    'that stage. From above, guests will see the story of the UAE\'s '
    'creativity and progress unfold beneath them: Louvre Abu Dhabi, Zayed '
    'National Museum, and Guggenheim Abu Dhabi, all within one sweeping '
    'view. Hero Experiences Group brings the product, the team, and the '
    'operational expertise. We are ready to deploy, operate, and deliver '
    'immediate value."',
    italic=True, color=MUTED, size=10)

# ── SAVE ─────────────────────────────────────────────────────────────
output_path = 'd:/Proposal-pitch/Original-vs-Current-Comparison.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
